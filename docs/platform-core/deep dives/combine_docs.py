# combine_docx.py
import os
import glob
from docx import Document

def combine_files(path):
    # Combine all .docx files in the given directory
    combined = Document()

    for filename in glob.glob(os.path.join(path, '*.docx')):
        doc = Document(filename)

        # Append content from each document to the combined document
        for paragraph in doc.paragraphs:
            combined.add_paragraph(paragraph.text)

    # Save the combined document as a new file
    combined.save('combined.docx')

if __name__ == "__main__":
    path = '/Users/sdjaeb/Downloads/drive-download-20250614T032554Z-1-001'
    combine_files(path)